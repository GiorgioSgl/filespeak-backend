from fastapi import UploadFile, File, HTTPException, APIRouter
from fastapi.responses import PlainTextResponse, HTMLResponse
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from decouple import config
from openai import OpenAI
from typing import List
from ..my_chromadb import collection
import tiktoken
import re

# Set up the API Key
client = OpenAI(api_key=config("OPENAI_API_KEY"))

router = APIRouter()

def chunk_text(text, max_tokens=8000):
    enc = tiktoken.encoding_for_model("text-embedding-3-small")
    tokens = enc.encode(text)
    chunks = []

    for i in range(0, len(tokens), max_tokens):
        chunk_tokens = tokens[i:i + max_tokens]
        chunk_text = enc.decode(chunk_tokens)
        chunks.append(chunk_text)

    return chunks

@router.post("/upload_multiple", response_class=PlainTextResponse, tags=["documents"])
async def upload_files(files: List[UploadFile] = File(...)):
    responses = []

    for file in files:
        filename = file.filename.lower()
        content = ""

        try:
            if filename.endswith(".txt") or filename.endswith(".html"):
                content = (await file.read()).decode("utf-8")

            elif filename.endswith(".pdf"):
                reader = PdfReader(BytesIO(await file.read()))
                content = "\n".join(page.extract_text() or "" for page in reader.pages)

            elif filename.endswith(".docx"):
                doc = Document(BytesIO(await file.read()))
                content = "\n".join([p.text for p in doc.paragraphs])

            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file format: {file.filename}")

            print(f"\n--- Processing {file.filename} ---")

            chunks = chunk_text(content, max_tokens=4000)

            if len(chunks) == 1:
                collection.add(
                    documents=[chunks[0]],
                    metadatas=[{"source": file.filename}],
                    ids=[file.filename]
                )
            else:
                collection.add(
                    documents=chunks,
                    metadatas=[{"source": file.filename} for _ in chunks],
                    ids=[f"{file.filename}_{i}" for i in range(len(chunks))]
                )            

            responses.append(f"{file.filename} processed successfully.")

        except Exception as e:
            responses.append(f"{file.filename} could not be processed: {str(e)}")

    return "\n".join(responses)


@router.post("/upload", response_class=PlainTextResponse, tags=["documents"])
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = ""

    if filename.endswith(".txt") or filename.endswith(".html"):
        content = (await file.read()).decode("utf-8")

    elif filename.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(await file.read()))
            content = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading PDF: {e}")

    elif filename.endswith(".docx"):
        try:
            doc = Document(BytesIO(await file.read()))
            content = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading DOCX: {e}")

    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")

    print(f"\n--- Content of {file.filename} ---\n{content}\n")

    collection.add(
        documents=[content],
        metadatas=[{"source": file.filename}],
        ids=[file.filename]
    )

    return f"File {file.filename} received and processed."


@router.get("/query", tags=["documents"])
async def query_file(query: str):
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    results = collection.query(
        query_texts=[query],
        n_results=1
    )

    if not results or not results['documents']:
        return "No results found."

    return results["ids"]


@router.get("/ls", tags=["documents"])
async def list_documents():
    try:
        data = collection.get()
        return data.get("ids", [])  # Return a list directly
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")


@router.post("/reset", tags=["documents"])
async def reset_collection():
    try:
        data = collection.get()
        ids = data.get("ids", [])
        if not ids:
            return {"message": "The collection is already empty."}
        collection.delete(ids=ids)
        return {"message": f"{len(ids)} documents deleted."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error resetting collection: {str(e)}")
