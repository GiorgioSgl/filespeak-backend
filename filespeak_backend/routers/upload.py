from fastapi import UploadFile, File, HTTPException, APIRouter
from fastapi.responses import PlainTextResponse, HTMLResponse
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from decouple import config
from openai import OpenAI
from typing import List
from ..my_chromadb import collection
# Configurar la API Key

client = OpenAI(api_key=config("OPENAI_API_KEY"))

router = APIRouter()


@router.post("/upload_multiple", response_class=PlainTextResponse, tags=["upload"])
async def upload_files(files: List[UploadFile] = File(...)):
    responses = []

    for file in files:
        filename = file.filename.lower()
        content = ""

        try:
            if filename.endswith(".txt") or filename.endswith(".html"):
                content = (await file.read()).decode("utf-8")

            elif filename.endswith(".pdf"):
                reader = PdfReader(BytesIO(await file.read()))
                content = "\n".join(page.extract_text() or "" for page in reader.pages)

            elif filename.endswith(".docx"):
                doc = Document(BytesIO(await file.read()))
                content = "\n".join([p.text for p in doc.paragraphs])

            else:
                raise HTTPException(status_code=400, detail=f"Formato de archivo no soportado: {file.filename}")

            print(f"\n--- Contenido de {file.filename} ---\n{content}\n")

            collection.add(
                documents=[content],
                metadatas=[{"source": file.filename}],
                ids=[file.filename]
            )

            responses.append(f"{file.filename} procesado correctamente.")

        except Exception as e:
            responses.append(f"{file.filename} no se pudo procesar: {str(e)}")

    return "\n".join(responses)


@router.post("/upload", response_class=PlainTextResponse, tags=["upload"])
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = ""

    if filename.endswith(".txt") or filename.endswith(".html"):
        content = (await file.read()).decode("utf-8")

    elif filename.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(await file.read()))
            content = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error leyendo PDF: {e}")

    elif filename.endswith(".docx"):
        try:
            doc = Document(BytesIO(await file.read()))
            content = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error leyendo DOCX: {e}")

    else:
        raise HTTPException(status_code=400, detail="Formato de archivo no soportado.")

    print(f"\n--- Contenido de {file.filename} ---\n{content}\n")

    collection.add(
        documents=[content],
        metadatas=[{"source": file.filename}],
        ids=[file.filename]
    )

    return f"Archivo {file.filename} recibido y procesado."

@router.get("/query", tags=["upload"])
async def query_file(query: str):
    if not query:
        raise HTTPException(status_code=400, detail="La consulta no puede estar vacía.")

    results = collection.query(
        query_texts=[query],
        n_results=1
    )

    if not results or not results['documents']:
        return "No se encontraron resultados."

    return results["ids"]


@router.get("/ls", tags=["upload"])
async def list_documents():
    try:
        data = collection.get()
        return data.get("ids", [])  # Devuelve una lista directamente
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")


@router.post("/reset", tags=["upload"])
async def reset_collection():
    try:
        data = collection.get()
        ids = data.get("ids", [])
        if not ids:
            return {"message": "La colección ya está vacía."}
        collection.delete(ids=ids)
        return {"message": f"Se eliminaron {len(ids)} documentos."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al resetear la colección: {str(e)}")
    